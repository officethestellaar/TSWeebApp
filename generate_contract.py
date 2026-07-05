#!/usr/bin/env python3
"""Generate a formal Developer Contract (.docx)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ─── Page Setup ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# ─── Color Palette ────────────────────────────────────────────────────────────
GOLD        = RGBColor(0x9B, 0x7B, 0x3C)
NAVY        = RGBColor(0x1B, 0x2A, 0x4A)
DARK_TEXT   = RGBColor(0x1A, 0x1A, 0x1A)
MEDIUM_TEXT = RGBColor(0x33, 0x33, 0x33)
LINE_COLOR  = RGBColor(0x9B, 0x7B, 0x3C)

# ─── Helper functions ─────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(color=LINE_COLOR, width_pt=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{width_pt * 8}" w:space="1" w:color="{color}" />'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_gold_line():
    add_horizontal_line(GOLD, 2)

def add_fancy_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 60)
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else GOLD
        run.font.name = 'Times New Roman'
    h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(8)
    return h

def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p

def body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.font.color.rgb = DARK_TEXT
    return p

def clause(num, text, indent_level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * indent_level)
    run = p.add_run(f'{num}.\t{text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_TEXT
    return p

def subclause(num, text, indent_level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.27 * indent_level)
    run = p.add_run(f'({num})\t{text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = MEDIUM_TEXT
    return p

def bullet(text, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = MEDIUM_TEXT
    return p

def field_line(label, width_inches=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run2 = p.add_run('_' * int(width_inches * 8))
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run2.font.size = Pt(12)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()  # spacer
doc.add_paragraph()
doc.add_paragraph()

# Gold emblem line
add_gold_line()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(30)
run = title.add_run('DEVELOPMENT SERVICES AGREEMENT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = NAVY
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('(Software Development & Technology Consulting)')
run.font.size = Pt(13)
run.font.color.rgb = GOLD
run.font.name = 'Times New Roman'
run.italic = True

add_gold_line()

doc.add_paragraph()

# Contract reference
ref = doc.add_paragraph()
ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ref.add_run('Agreement No: _______________')
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_TEXT
run.font.name = 'Times New Roman'

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_after = Pt(4)
run = date_p.add_run(f'Date: {datetime.date.today().strftime("%d %B %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_TEXT
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

# Parties
parties_heading = doc.add_paragraph()
parties_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = parties_heading.add_run('BETWEEN')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = NAVY
run.font.name = 'Times New Roman'

doc.add_paragraph()

# Party 1
party1_box = doc.add_table(rows=1, cols=1)
party1_box.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = party1_box.cell(0, 0)
cell.text = ''
set_cell_shading(cell, 'F5F0E8')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
run = p.add_run('THE STELLAAR CLUB')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = NAVY
run.font.name = 'Times New Roman'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('(hereinafter referred to as the "Client")')
run2.italic = True
run2.font.size = Pt(11)
run2.font.color.rgb = MEDIUM_TEXT
run2.font.name = 'Times New Roman'

doc.add_paragraph()
and_p = doc.add_paragraph()
and_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = and_p.add_run('— AND —')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = GOLD
run.font.name = 'Times New Roman'

doc.add_paragraph()

# Party 2
party2_box = doc.add_table(rows=1, cols=1)
party2_box.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = party2_box.cell(0, 0)
cell.text = ''
set_cell_shading(cell, 'F5F0E8')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
run = p.add_run('[DEVELOPER NAME]')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = NAVY
run.font.name = 'Times New Roman'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('(hereinafter referred to as the "Developer")')
run2.italic = True
run2.font.size = Pt(11)
run2.font.color.rgb = MEDIUM_TEXT
run2.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

collectively = doc.add_paragraph()
collectively.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = collectively.add_run('(Collectively referred to as the "Parties" and individually as a "Party")')
run.font.size = Pt(11)
run.font.color.rgb = GOLD
run.font.name = 'Times New Roman'
run.italic = True

doc.add_paragraph()

# Page break
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# RECITALS
# ═══════════════════════════════════════════════════════════════════════════════

heading('RECITALS', level=1)

recitals = [
    "WHEREAS the Client is engaged in the business of operating a premium membership-based club and desires to develop, maintain, and enhance its digital technology infrastructure, including but not limited to web applications, mobile interfaces, backend systems, and associated software components (collectively, the \"Project\");",
    "WHEREAS the Developer represents that it possesses the requisite skill, expertise, and technical capability to perform the development services contemplated under this Agreement;",
    "WHEREAS the Parties wish to set forth the terms and conditions under which the Developer shall provide development services to the Client.",
    "NOW, THEREFORE, in consideration of the mutual covenants, promises, and agreements contained herein, the Parties agree as follows:"
]

for r in recitals:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(r)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_TEXT

add_gold_line()

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 1 — DEFINITIONS
# ═══════════════════════════════════════════════════════════════════════════════

heading('ARTICLE 1: DEFINITIONS', level=1)

body("For the purposes of this Agreement, the following terms shall have the meanings ascribed to them below:")

definitions = [
    ("1.1", "\"Agreement\"", "means this Development Services Agreement, including all schedules, exhibits, and addenda attached hereto."),
    ("1.2", "\"Confidential Information\"", "means any and all technical and business information disclosed by one Party to the other, including but not limited to source code, algorithms, database schemas, API keys, business processes, financial data, member information, and trade secrets."),
    ("1.3", "\"Deliverables\"", "means the software, code, documentation, designs, and other materials created by the Developer under this Agreement."),
    ("1.4", "\"Effective Date\"", "means the date of last signature of this Agreement."),
    ("1.5", "\"Intellectual Property Rights\"", "means all patents, copyrights, trademarks, trade secrets, and any other proprietary rights recognized in any jurisdiction."),
    ("1.6", "\"Project\"", "means the software development and technology services described in Schedule A attached hereto."),
    ("1.7", "\"Statement of Work\" or \"SOW\"", "means a written document describing a specific scope of work, deliverables, timeline, and fees, which once signed by both Parties becomes part of this Agreement."),
    ("1.8", "\"Third-Party Materials\"", "means any software, libraries, tools, or other materials owned by a third party that are incorporated into or used in connection with the Deliverables."),
]

for num, term, defn in definitions:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.27)
    run1 = p.add_run(f'{num} ')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run1.font.color.rgb = DARK_TEXT
    run2 = p.add_run(term)
    run2.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.font.color.rgb = NAVY
    run3 = p.add_run(f' {defn}')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)
    run3.font.color.rgb = DARK_TEXT

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 2 — SCOPE OF WORK
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 2: SCOPE OF WORK', level=1)

clause("2.1", "The Developer agrees to perform the development services described in the Statement(s) of Work attached hereto as Schedule A and any subsequent SOWs executed by the Parties.")

clause("2.2", "Each SOW shall include, at a minimum:")

subclause("a", "A detailed description of the services and deliverables to be provided;")
subclause("b", "Project milestones, timelines, and delivery schedule;")
subclause("c", "The applicable fees, payment schedule, and billing structure;")
subclause("d", "Acceptance criteria and testing procedures;")
subclause("e", "Any specific technology stack, platforms, or frameworks to be used.")

clause("2.3", "The Developer shall perform the services in a professional and workmanlike manner, conforming to industry standards for software development.")

clause("2.4", "The Client shall provide timely access to necessary systems, credentials, documentation, and personnel reasonably required for the Developer to perform the services.")

clause("2.5", "Any changes to the scope of work shall be documented in a written change order signed by both Parties. No oral modifications shall be binding.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 3 — COMPENSATION & PAYMENT
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 3: COMPENSATION & PAYMENT', level=1)

clause("3.1", "The Client shall pay the Developer the fees set forth in the applicable SOW. Unless otherwise specified, fees are denominated in Indian Rupees (INR).")

clause("3.2", "Payment Terms:")

subclause("a", "Invoices shall be rendered by the Developer upon completion of each milestone or on a monthly basis, as specified in the SOW;")
subclause("b", "All invoices are due within fifteen (15) days of receipt;")
subclause("c", "Late payments shall accrue interest at the rate of 1.5% per month on the outstanding balance;")
subclause("d", "The Client shall reimburse the Developer for all pre-approved out-of-pocket expenses incurred in connection with the Project.")

clause("3.3", "All fees are exclusive of applicable taxes, including GST, which shall be borne by the Client.")

clause("3.4", "The Developer reserves the right to suspend performance of services if any invoice remains unpaid for more than thirty (30) days beyond the due date, upon providing five (5) business days' written notice.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 4 — INTELLECTUAL PROPERTY
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 4: INTELLECTUAL PROPERTY RIGHTS', level=1)

clause("4.1", "Subject to full payment of all fees due under this Agreement, the Developer hereby assigns to the Client all right, title, and interest in and to the Deliverables, including all Intellectual Property Rights embodied therein.")

clause("4.2", "The Developer retains the right to:")

subclause("a", "Use general skills, knowledge, and expertise gained through performing the services;")
subclause("b", "Utilize any pre-existing tools, libraries, frameworks, or methodologies owned by the Developer prior to this Agreement (\"Developer Pre-Existing IP\"), provided that the Client is granted a perpetual, royalty-free, non-exclusive license to use such Pre-Existing IP as incorporated into the Deliverables;")
subclause("c", "Maintain a portfolio of work, including anonymous or de-identified descriptions of the Project.")

clause("4.3", "Third-Party Materials incorporated into the Deliverables shall be licensed to the Client under the applicable open-source or commercial license terms governing such materials.")

clause("4.4", "The Developer warrants that the Deliverables will not infringe upon the intellectual property rights of any third party.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 5 — CONFIDENTIALITY
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 5: CONFIDENTIALITY', level=1)

clause("5.1", "Each Party (\"Receiving Party\") agrees to hold in strict confidence all Confidential Information disclosed by the other Party (\"Disclosing Party\").")

clause("5.2", "The Receiving Party shall:")

subclause("a", "Not disclose Confidential Information to any third party without the prior written consent of the Disclosing Party;")
subclause("b", "Use Confidential Information solely for the purpose of performing obligations under this Agreement;")
subclause("c", "Protect Confidential Information using the same degree of care used to protect its own confidential information, but in no event less than reasonable care.")

clause("5.3", "The obligations of confidentiality shall not apply to information that:")

subclause("a", "Is or becomes publicly available through no fault of the Receiving Party;")
subclause("b", "Was lawfully in the Receiving Party's possession prior to disclosure;")
subclause("c", "Is independently developed by the Receiving Party without use of Confidential Information;")
subclause("d", "Is required to be disclosed by applicable law or court order.")

clause("5.4", "The obligations of this Article 5 shall survive termination of this Agreement for a period of five (5) years.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 6 — NON-COMPETE & NON-SOLICITATION
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 6: NON-COMPETE & NON-SOLICITATION', level=1)

clause("6.1", "During the term of this Agreement and for a period of twelve (12) months thereafter, the Developer shall not, directly or indirectly, solicit, induce, or encourage any employee, contractor, or consultant of the Client to terminate their relationship with the Client.")

clause("6.2", "The Developer shall not, during the term of this Agreement, provide development services to any entity that is a direct competitor of the Client in the premium club membership market within the territory of India, without the prior written consent of the Client.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 7 — TERM & TERMINATION
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 7: TERM & TERMINATION', level=1)

clause("7.1", "This Agreement shall commence on the Effective Date and shall continue until terminated as provided herein.")

clause("7.2", "Either Party may terminate this Agreement upon thirty (30) days' prior written notice to the other Party.")

clause("7.3", "Either Party may terminate this Agreement immediately upon written notice if:")

subclause("a", "The other Party commits a material breach and fails to cure such breach within fifteen (15) days of receiving written notice thereof;")
subclause("b", "The other Party becomes insolvent, files for bankruptcy, or ceases operations.")

clause("7.4", "Upon termination:")

subclause("a", "The Developer shall deliver all completed and in-progress Deliverables to the Client;")
subclause("b", "The Client shall pay all fees due for services rendered up to the date of termination;")
subclause("c", "Articles 4 (Intellectual Property), 5 (Confidentiality), 6 (Non-Compete), 9 (Indemnification), and 10 (Limitation of Liability) shall survive termination.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 8 — REPRESENTATIONS & WARRANTIES
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 8: REPRESENTATIONS & WARRANTIES', level=1)

clause("8.1", "Each Party represents and warrants that it has the full right, power, and authority to enter into this Agreement.")

clause("8.2", "The Developer warrants that:")

subclause("a", "The Deliverables will conform to the specifications set forth in the applicable SOW;")
subclause("b", "The Deliverables will be free from material defects in design, coding, and implementation for a period of ninety (90) days from delivery (\"Warranty Period\");")
subclause("c", "The services will be performed in a professional manner consistent with industry standards.")

clause("8.3", "During the Warranty Period, the Developer shall, at no additional cost, correct any material non-conformities or defects in the Deliverables promptly upon notification by the Client.")

clause("8.4", "EXCEPT AS EXPRESSLY SET FORTH IN THIS AGREEMENT, THE DEVELOPER MAKES NO OTHER WARRANTIES, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 9 — INDEMNIFICATION
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 9: INDEMNIFICATION', level=1)

clause("9.1", "The Developer agrees to indemnify, defend, and hold harmless the Client from and against any and all losses, damages, costs, and expenses (including reasonable legal fees) arising out of any third-party claim that the Deliverables infringe upon the intellectual property rights of such third party.")

clause("9.2", "The Client agrees to indemnify, defend, and hold harmless the Developer from and against any claims arising out of:")

subclause("a", "The Client's misuse of the Deliverables;")
subclause("b", "Any content, data, or materials provided by the Client that violate applicable law or third-party rights.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 10 — LIMITATION OF LIABILITY
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 10: LIMITATION OF LIABILITY', level=1)

clause("10.1", "TO THE MAXIMUM EXTENT PERMITTED BY APPLICABLE LAW, NEITHER PARTY SHALL BE LIABLE TO THE OTHER FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING BUT NOT LIMITED TO LOSS OF PROFITS, LOSS OF DATA, OR BUSINESS INTERRUPTION.")

clause("10.2", "The total cumulative liability of either Party under this Agreement shall not exceed the total fees paid or payable by the Client to the Developer during the twelve (12) months preceding the event giving rise to the claim.")

clause("10.3", "The limitations in this Article 10 shall not apply to:")

subclause("a", "Claims arising from a Party's breach of confidentiality obligations;")
subclause("b", "Claims arising from a Party's infringement of intellectual property rights;")
subclause("c", "Claims arising from fraud, gross negligence, or willful misconduct.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 11 — INSURANCE
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 11: INSURANCE', level=1)

clause("11.1", "The Developer shall maintain, at its own expense, the following insurance coverage for the duration of this Agreement:")

subclause("a", "Professional Liability / Errors & Omissions Insurance: minimum INR 10,00,000 per claim;")
subclause("b", "Commercial General Liability Insurance: minimum INR 5,00,000 per occurrence;")
subclause("c", "Cyber Liability Insurance: minimum INR 5,00,000 per occurrence.")

clause("11.2", "Upon the Client's request, the Developer shall provide certificates of insurance evidencing the coverage required herein.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 12 — DATA PROTECTION & SECURITY
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 12: DATA PROTECTION & SECURITY', level=1)

clause("12.1", "The Developer acknowledges that it may have access to personal data of the Client's members, employees, and other individuals. The Developer agrees to:")

subclause("a", "Process all personal data in compliance with applicable data protection laws, including the Information Technology Act, 2000, and the Digital Personal Data Protection Act, 2023;")
subclause("b", "Implement appropriate technical and organizational measures to protect personal data against unauthorized access, loss, or destruction;")
subclause("c", "Not use, share, or disclose personal data for any purpose other than performing the services under this Agreement;")
subclause("d", "Promptly notify the Client of any data breach or security incident affecting the Client's data.")

clause("12.2", "The Developer shall not store, process, or transmit the Client's data from servers located outside India without the Client's prior written consent.")

clause("12.3", "Upon termination of this Agreement, the Developer shall securely delete or return all personal data and Confidential Information in its possession within fifteen (15) business days.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 13 — INDEPENDENT CONTRACTOR
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 13: INDEPENDENT CONTRACTOR STATUS', level=1)

clause("13.1", "The Developer is an independent contractor and nothing in this Agreement shall be construed to create an employer-employee, partnership, joint venture, or agency relationship between the Parties.")

clause("13.2", "The Developer shall be solely responsible for:")

subclause("a", "Payment of all taxes, duties, and social security contributions;")
subclause("b", "Obtaining and maintaining any necessary licenses, permits, or registrations;")
subclause("c", "Providing all tools, equipment, and resources necessary to perform the services, unless otherwise agreed.")

clause("13.3", "The Developer is not authorized to bind the Client to any contract or obligation, and shall not represent itself as an agent or employee of the Client.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 14 — GOVERNING LAW & DISPUTE RESOLUTION
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 14: GOVERNING LAW & DISPUTE RESOLUTION', level=1)

clause("14.1", "This Agreement shall be governed by and construed in accordance with the laws of the Republic of India.")

clause("14.2", "Any dispute arising out of or relating to this Agreement shall be resolved through the following escalating mechanism:")

subclause("a", "Negotiation: The Parties shall first attempt to resolve the dispute through good-faith negotiations for a period of fourteen (14) days;")
subclause("b", "Mediation: If the dispute is not resolved through negotiation, the Parties shall submit the dispute to mediation before a mutually agreed mediator in Mumbai;")
subclause("c", "Arbitration: If the dispute is not resolved through mediation, it shall be finally settled by arbitration in accordance with the Arbitration and Conciliation Act, 1996. The arbitration shall be conducted in Mumbai by a sole arbitrator appointed by mutual agreement. The language of arbitration shall be English.")

clause("14.3", "The prevailing Party in any dispute resolution proceeding shall be entitled to recover its reasonable legal fees and costs from the non-prevailing Party.")

clause("14.4", "Notwithstanding the foregoing, either Party may seek injunctive or other equitable relief from a court of competent jurisdiction to prevent or restrain a breach of confidentiality or intellectual property rights.")

# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE 15 — MISCELLANEOUS
# ═══════════════════════════════════════════════════════════════════════════════

add_gold_line()
heading('ARTICLE 15: MISCELLANEOUS', level=1)

clause("15.1", "Entire Agreement: This Agreement constitutes the entire agreement between the Parties with respect to the subject matter hereof and supersedes all prior agreements, understandings, and communications.")

clause("15.2", "Amendments: No modification or amendment to this Agreement shall be effective unless in writing and signed by both Parties.")

clause("15.3", "Waiver: The failure of either Party to enforce any provision of this Agreement shall not be deemed a waiver of such provision or any subsequent breach.")

clause("15.4", "Severability: If any provision of this Agreement is held to be invalid or unenforceable, the remaining provisions shall continue in full force and effect.")

clause("15.5", "Notices: All notices under this Agreement shall be in writing and delivered by email, registered post, or courier to the addresses set forth below.")

clause("15.6", "Force Majeure: Neither Party shall be liable for any delay or failure in performance caused by events beyond its reasonable control, including but not limited to acts of God, war, terrorism, strikes, and governmental actions.")

clause("15.7", "Assignment: Neither Party may assign this Agreement without the prior written consent of the other Party, except that the Developer may assign this Agreement to a wholly-owned subsidiary or successor entity.")

clause("15.8", "Counterparts: This Agreement may be executed in two or more counterparts, each of which shall be deemed an original and all of which together shall constitute one instrument.")

# ═══════════════════════════════════════════════════════════════════════════════
# SCHEDULE A — STATEMENT OF WORK
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_gold_line()
heading('SCHEDULE A: STATEMENT OF WORK', level=1)

add_fancy_line()

subheading('A. PROJECT OVERVIEW')

field_line('Project Name')
field_line('Project Reference')
field_line('Effective Date')

doc.add_paragraph()
subheading('B. SCOPE OF SERVICES & DELIVERABLES')

body("The Developer shall perform the following services and deliver the following deliverables:")

for i in range(5):
    bullet('_______________________________________________________________________________')

doc.add_paragraph()
subheading('C. TECHNOLOGY STACK')

body("The project shall be developed using the following technology stack:")

field_line('Frontend Framework')
field_line('Backend Framework')
field_line('Database')
field_line('Hosting / Infrastructure')
field_line('Third-Party Integrations')

doc.add_paragraph()
subheading('D. MILESTONES & TIMELINE')

# Milestones table
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Milestone', 'Description', 'Delivery Date', 'Fee (INR)']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    set_cell_shading(cell, '1B2A4A')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

milestones_data = [
    ['Milestone 1', '', '', ''],
    ['Milestone 2', '', '', ''],
    ['Milestone 3', '', '', ''],
    ['Milestone 4', '', '', ''],
    ['Final Delivery', '', '', ''],
]

for r_idx, row_data in enumerate(milestones_data, 1):
    for c_idx, val in enumerate(row_data):
        cell = table.rows[r_idx].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        if c_idx == 0:
            run.bold = True

doc.add_paragraph()

# Total
total_table = doc.add_table(rows=1, cols=4)
total_table.style = 'Table Grid'
total_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i in range(4):
    cell = total_table.rows[0].cells[i]
    cell.text = ''
    set_cell_shading(cell, 'F5F0E8')
    if i < 3:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run('TOTAL' if i == 2 else '')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    else:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('₹ _______________')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc.add_paragraph()
subheading('E. PAYMENT SCHEDULE')

bullet('__________________% upon signing of this Agreement')
bullet('__________________% upon completion of Milestone 1')
bullet('__________________% upon completion of Milestone 2')
bullet('__________________% upon completion of Milestone 3')
bullet('__________________% upon Final Delivery & Acceptance')

doc.add_paragraph()
subheading('F. ACCEPTANCE CRITERIA')

bullet('All Deliverables pass functional testing against the approved specifications')
bullet('Code is delivered with appropriate documentation and deployment instructions')
bullet('No critical or high-severity bugs are present at the time of delivery')
bullet('All Third-Party Materials are properly licensed and attributed')

doc.add_paragraph()
subheading('G. ADDITIONAL TERMS')

body('_______________________________________________________________________________')
body('_______________________________________________________________________________')
body('_______________________________________________________________________________')

# ═══════════════════════════════════════════════════════════════════════════════
# SIGNATURE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_gold_line()
heading('IN WITNESS WHEREOF', level=1)

body("The Parties have executed this Agreement as of the Effective Date.")

doc.add_paragraph()
doc.add_paragraph()

# Signature table
sig_table = doc.add_table(rows=5, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Column widths
for row in sig_table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = Inches(3) if i == 0 else Inches(3)

# Left - Client
cell = sig_table.cell(0, 0)
cell.text = ''
set_cell_shading(cell, 'F5F0E8')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('THE STELLAAR CLUB')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = NAVY
run.font.name = 'Times New Roman'

cell = sig_table.cell(1, 0)
cell.text = ''
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Signature:')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('_' * 30)
run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

cell = sig_table.cell(2, 0)
cell.text = ''
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Name:')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('_' * 30)

cell = sig_table.cell(3, 0)
cell.text = ''
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Title:')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('_' * 30)

cell = sig_table.cell(4, 0)
cell.text = ''
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Date:')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('_' * 30)

# Right - Developer
cell = sig_table.cell(0, 1)
cell.text = ''
set_cell_shading(cell, 'F5F0E8')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('[DEVELOPER NAME]')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = NAVY
run.font.name = 'Times New Roman'

cell = sig_table.cell(1, 1)
cell.text = ''
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Signature:')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('_' * 30)

cell = sig_table.cell(2, 1)
cell.text = ''
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Name:')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('_' * 30)

cell = sig_table.cell(3, 1)
cell.text = ''
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Title:')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('_' * 30)

cell = sig_table.cell(4, 1)
cell.text = ''
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Date:')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('_' * 30)

doc.add_paragraph()
add_gold_line()
doc.add_paragraph()

# Footer
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('THE STELLAAR CLUB — Development Services Agreement')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Times New Roman'
run.italic = True

# ─── Save ─────────────────────────────────────────────────────────────────────
output_path = 'DEVELOPMENT_SERVICES_AGREEMENT.docx'
doc.save(output_path)
print(f'Generated: {output_path}')
